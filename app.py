import streamlit as st
import re
import fitz
import io
import docx
import uuid
import glob
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import requests
import json
from datetime import datetime
from io import BytesIO
import os
import time
import pickle
import hashlib

# Configuration
st.set_page_config(
    page_title="MBE Document Assistant - RAG Chatbot",
    page_icon="🎓",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        color: white;
        margin-bottom: 2rem;
    }
    .stat-box {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .chunk-display {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #667eea;
        white-space: pre-wrap;
        font-family: 'Courier New', monospace;
        line-height: 1.6;
    }
    .answer-box {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    .citation-tag {
        background: #667eea;
        color: white;
        padding: 0.2rem 0.5rem;
        border-radius: 5px;
        font-size: 0.85em;
        margin-right: 0.5rem;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .user-message {
        background: #e3f2fd;
        margin-left: 2rem;
    }
    .assistant-message {
        background: #f3e5f5;
        margin-right: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'processed' not in st.session_state:
    st.session_state.processed = False
    st.session_state.files_data = {}
    st.session_state.collection = None
    st.session_state.messages = []  # Chat history
    st.session_state.current_context = []  # للـ conversational flow

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "llama-3.3-70b-versatile"
PDF_PASSWORD = "mbe2025"
DOCS_FOLDER = "/mount/src/test/documents"
CACHE_FOLDER = os.getenv("CACHE_FOLDER", "./cache")

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(CACHE_FOLDER, exist_ok=True)

if not GROQ_API_KEY:
    st.error("⚠️ GROQ_API_KEY not found in environment variables!")

# Helper Functions
def get_file_hash(filepath):
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def load_cache(cache_key):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        except:
            return None
    return None

def save_cache(cache_key, data):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    try:
        with open(cache_file, 'wb') as f:
            pickle.dump(data, f)
    except Exception as e:
        st.warning(f"⚠️ Could not save cache: {str(e)}")

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
    return text.strip()

def structure_text_into_paragraphs(text):
    if not text or not text.strip():
        return ""
   
    text = clean_text(text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
   
    if not lines:
        return ""
   
    paragraphs = []
    current_paragraph = []
   
    for i, line in enumerate(lines):
        words_in_line = line.split()
       
        if len(words_in_line) < 3 and not (line[0].isupper() or re.match(r'^[\d]+[\.\):]', line)):
            continue
       
        is_heading = (
            (line.isupper() and len(words_in_line) <= 10) or
            (len(words_in_line) <= 6 and line[0].isupper() and line.endswith(':'))
        )
       
        if is_heading:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f"\n🔹 {line}\n")
            continue
       
        is_list_item = re.match(r'^[\d]+[\.\)]\s', line) or re.match(r'^[•\-\*]\s', line)
       
        if is_list_item:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f" {line}")
            continue
       
        current_paragraph.append(line)
       
        ends_with_punctuation = line.endswith(('.', '!', '?', '؟', '!', '。'))
        next_is_new_section = False
       
        if i < len(lines) - 1:
            next_line = lines[i + 1]
            next_words = next_line.split()
            next_is_new_section = (
                re.match(r'^[\d]+[\.\)]\s', next_line) or
                re.match(r'^[•\-\*]\s', next_line) or
                (len(next_words) <= 6 and next_line[0].isupper()) or
                next_line.isupper()
            )
       
        is_last_line = (i == len(lines) - 1)
       
        if (ends_with_punctuation or next_is_new_section or is_last_line):
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraph_text = re.sub(r'([.,!?;:])\s*([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
   
    if paragraphs:
        structured_text = ""
        for para in paragraphs:
            if para.startswith('\n🔹'):
                structured_text += para
            elif para.startswith(' '):
                structured_text += para + "\n"
            else:
                structured_text += para + "\n\n"
        return structured_text.strip()
   
    return text

def create_smart_chunks(text, chunk_size=1000, overlap=200, page_num=None, source_file=None, is_table=False, table_num=None):
    """Enhanced chunking with metadata"""
    words = text.split()
    chunks = []
    
    # Clean metadata - ChromaDB doesn't accept None values
    metadata = {
        'page': str(page_num) if page_num is not None else "N/A",
        'source': source_file if source_file else "Unknown",
        'is_table': str(is_table),
        'table_number': str(table_num) if table_num is not None else "N/A"
    }
   
    if len(words) <= chunk_size:
        if text.strip():
            return [{
                'content': text,
                'metadata': metadata
            }]
        return []
   
    for i in range(0, len(words), chunk_size - overlap):
        chunk_words = words[i:i + chunk_size]
        chunk = " ".join(chunk_words)
        if len(chunk.split()) >= 30:
            chunks.append({
                'content': chunk,
                'metadata': metadata.copy()
            })
   
    return chunks

def format_table_as_structured_text(extracted_table, table_number=None):
    if not extracted_table or len(extracted_table) == 0:
        return ""
   
    headers = [str(cell).strip() if cell else f"Column_{i+1}" for i, cell in enumerate(extracted_table[0])]
    headers = [clean_text(h) for h in headers]
   
    text = f"\n📊 Table {table_number or ''}\n\n"
    if headers:
        text += "| " + " | ".join(headers) + " |\n"
        text += "| " + " --- |" * len(headers) + " |\n"
    
    row_count = 0
    for row in extracted_table[1:]:
        cells = [str(cell).strip() if cell else "" for cell in row]
        cells = [clean_text(c) for c in cells]
        if any(cells):
            text += "| " + " | ".join(cells) + " |\n"
            row_count += 1
    
    text += f"\n**Summary**: {row_count} data rows, {len(headers)} columns.\n"
    return text

def extract_pdf_detailed(filepath):
    try:
        doc = fitz.open(filepath)
        if doc.is_encrypted:
            if not doc.authenticate(PDF_PASSWORD):
                doc.close()
                return None, "❌ Invalid PDF password"
    except Exception as e:
        return None, f"❌ Error opening PDF: {str(e)}"
   
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': len(doc),
        'total_tables': 0,
        'pages_with_tables': [],
    }
   
    for page_num in range(len(doc)):
        page = doc[page_num]
        all_elements = []
       
        text_blocks = page.get_text("dict")["blocks"]
        for block in text_blocks:
            if block.get('type') == 0:
                y_pos = block.get('bbox', [0, 0, 0, 0])[1]
                text_content = ""
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        text_content += span.get('text', '') + ' '
                if text_content.strip():
                    structured_content = structure_text_into_paragraphs(text_content)
                    all_elements.append({
                        'type': 'text',
                        'y_position': y_pos,
                        'content': structured_content,
                        'page': page_num + 1
                    })
       
        tables = page.find_tables()
        if tables and len(tables.tables) > 0:
            file_info['pages_with_tables'].append(page_num + 1)
           
            for table_num, table in enumerate(tables.tables, 1):
                file_info['total_tables'] += 1
                extracted_table = table.extract()
                if extracted_table:
                    table_text = format_table_as_structured_text(extracted_table, file_info['total_tables'])
                    all_elements.append({
                        'type': 'table',
                        'y_position': table.bbox[1] if table.bbox else 0,
                        'content': table_text,
                        'page': page_num + 1,
                        'table_num': file_info['total_tables']
                    })
       
        all_elements.sort(key=lambda x: x['y_position'])
       
        # Add document title header at the start of each page
        page_text = f"\n# Document: {filename} - Official MBE Regulations\n\n"
        page_text += f"\n{'═' * 60}\n📄 Page {page_num + 1}\n{'═' * 60}\n\n"
        for element in all_elements:
            page_text += element['content'] + "\n\n"
       
        # Create chunks with metadata
        page_chunks = create_smart_chunks(
            page_text, 
            chunk_size=1500, 
            overlap=250,
            page_num=page_num + 1,
            source_file=filename,
            is_table=False
        )
        
        # Add table chunks separately with table metadata
        for element in all_elements:
            if element['type'] == 'table':
                table_chunks = create_smart_chunks(
                    element['content'],
                    chunk_size=2000,  # Larger size for tables
                    overlap=0,
                    page_num=element['page'],
                    source_file=filename,
                    is_table=True,
                    table_num=element.get('table_num')
                )
                file_info['chunks'].extend(table_chunks)
        
        file_info['chunks'].extend(page_chunks)
   
    doc.close()
    return file_info, None

def extract_docx_detailed(filepath):
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
   
    all_text = []
    table_counter = 0
   
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = clean_text(para.text)
                    if text:
                        structured = structure_text_into_paragraphs(text)
                        if structured:
                            all_text.append(structured)
                    break
       
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    file_info['total_tables'] += 1
                    table_counter += 1
                    table_text = format_table_as_structured_text(
                        [[cell.text for cell in row.cells] for row in table.rows],
                        table_counter
                    )
                    if table_text:
                        all_text.append(table_text)
                        # Add table as separate chunk
                        table_chunks = create_smart_chunks(
                            table_text,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=filename,
                            is_table=True,
                            table_num=table_counter
                        )
                        file_info['chunks'].extend(table_chunks)
                    break
   
    complete_text = "\n\n".join(all_text)
    text_chunks = create_smart_chunks(
        complete_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename
    )
    file_info['chunks'].extend(text_chunks)
   
    if file_info['total_tables'] > 0:
        file_info['pages_with_tables'] = [1]
   
    return file_info, None

def extract_txt_detailed(filepath):
    filename = os.path.basename(filepath)
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    structured_text = structure_text_into_paragraphs(text)
    chunks = create_smart_chunks(
        structured_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename
    )
    file_info = {
        'chunks': chunks,
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
    return file_info, None

def get_files_from_folder():
    supported_extensions = ['*.pdf', '*.docx', '*.doc', '*.txt']
    files = []
    for ext in supported_extensions:
        files.extend(glob.glob(os.path.join(DOCS_FOLDER, ext)))
    return files

def get_embedding_function():
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="intfloat/multilingual-e5-large"
    )

def build_conversational_prompt(query, chat_history):
    """Build context-aware prompt with chat history"""
    if not chat_history:
        return query
    
    # Last 3 exchanges for context
    recent_history = chat_history[-6:]  # 3 Q&A pairs
    context_lines = []
    
    for msg in recent_history:
        role = msg['role']
        content = msg['content'][:200]  # Limit length
        if role == 'user':
            context_lines.append(f"Previous Q: {content}")
        else:
            context_lines.append(f"Previous A: {content}")
    
    history_context = "\n".join(context_lines)
    return f"Conversation context:\n{history_context}\n\nCurrent question: {query}"

def answer_question_with_groq(query, relevant_chunks, chat_history=None):
    if not GROQ_API_KEY:
        return "❌ Please set GROQ_API_KEY in environment variables"
   
    # Build context with precise citations
    context_parts = []
    for i, chunk_data in enumerate(relevant_chunks[:10], 1):
        content = chunk_data['content']
        meta = chunk_data['metadata']
        
        # Handle string metadata from ChromaDB
        source = meta.get('source', 'Unknown')
        page = meta.get('page', 'N/A')
        is_table = meta.get('is_table', 'False')
        table_num = meta.get('table_number', 'N/A')
        
        citation = f"[Source {i}: {source}, Page {page}"
        if is_table == 'True' or is_table == True:
            citation += f", Table {table_num}"
        citation += "]"
        
        context_parts.append(f"{citation}\n{content}")
    
    context = "\n\n---\n\n".join(context_parts)
    
    # Build conversation history for follow-ups
    conversation_summary = ""
    if chat_history and len(chat_history) > 0:
        recent = chat_history[-6:]  # Last 3 Q&A pairs
        conv_lines = []
        for msg in recent:
            role = "User" if msg['role'] == 'user' else "Assistant"
            content_preview = msg['content'][:300]  # Longer preview
            conv_lines.append(f"{role}: {content_preview}")
        conversation_summary = "\n".join(conv_lines)
   
    data = {
        "model": GROQ_MODEL,
        "messages": [
            {
                "role": "system",
                "content": """You are a precise MBE Document Assistant at Hochschule Anhalt specializing in Biomedical Engineering regulations.

CRITICAL RULES:
1. Answer ONLY from provided sources OR previous conversation if it's a follow-up question
2. ALWAYS cite sources: [Source X, Page Y] or [Source X, Page Y, Table Z]
3. For follow-up questions like "summarize", "tell me more", "explain that", or "what about that":
   - Check the conversation history FIRST
   - Summarize or expand on your PREVIOUS answer
   - Don't search for new information if the question refers to what you just said
4. If user says "summarize that" or "summarize it": Condense your LAST answer (from conversation history)
5. If no relevant info in sources OR history: "No sufficient information in the available documents"
6. Use the SAME language as the question (English/German/Arabic)
7. Be CONCISE - short, direct answers unless asked to elaborate
8. For counting questions: Count precisely and list all items with citations

FOLLOW-UP DETECTION:
- "that", "it", "this", "summarize", "tell me more", "elaborate", "explain further" → Use conversation history
- New factual questions → Use sources

Remember: You're helping MBE students understand their program requirements clearly and accurately."""
            },
            {
                "role": "user",
                "content": f"""CONVERSATION HISTORY (use for follow-up questions):
{conversation_summary if conversation_summary else "No previous conversation"}

DOCUMENT SOURCES (use for new factual questions):
{context}

CURRENT QUESTION: {query}

Instructions: 
- If this is a follow-up (summarize/elaborate/that/it), answer from conversation history
- If this is a new question, answer from sources with citations
- Always be precise and cite your sources

ANSWER:"""
            }
        ],
        "temperature": 0.1,  # Slightly higher for better conversational flow
        "max_tokens": 2000,
    }
   
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json=data,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"❌ Error connecting to Groq: {str(e)}"

# Main UI
st.markdown("""
<div class="main-card">
    <h1 style='text-align: center; margin: 0;'>🎓 MBE Document Assistant</h1>
    <p style='text-align: center; margin-top: 10px;'>RAG Chatbot for Biomedical Engineering at Hochschule Anhalt</p>
</div>
""", unsafe_allow_html=True)

# Sidebar for document management
with st.sidebar:
    st.markdown("### 📚 Document Management")
    
    available_files = get_files_from_folder()
    if not available_files:
        st.warning(f"⚠️ No documents found")
        st.info(f"📁 Add files to: {DOCS_FOLDER}")
    else:
        st.success(f"✅ {len(available_files)} document(s)")
        with st.expander("📂 Files", expanded=False):
            for file in available_files:
                st.write(f"• {os.path.basename(file)}")
    
    st.markdown("---")
    
    if available_files and st.button("🚀 Process Documents", type="primary", use_container_width=True):
        with st.spinner("Processing..."):
            files_data = {}
            all_chunks = []
            all_metadata = []
           
            client = chromadb.Client()
            collection_name = f"docs_{uuid.uuid4().hex[:8]}"
            collection = client.create_collection(
                name=collection_name,
                embedding_function=get_embedding_function()
            )
           
            progress_bar = st.progress(0)
            status_text = st.empty()
           
            for idx, filepath in enumerate(available_files):
                filename = os.path.basename(filepath)
                file_ext = filename.split('.')[-1].lower()
               
                status_text.text(f"Processing: {filename}...")
               
                file_hash = get_file_hash(filepath)
                cache_key = f"{file_hash}_{file_ext}"
                cached_data = load_cache(cache_key)
               
                if cached_data:
                    st.info(f"📦 Cached: {filename}")
                    file_info = cached_data
                    error = None
                else:
                    if file_ext == 'pdf':
                        file_info, error = extract_pdf_detailed(filepath)
                    elif file_ext in ['docx', 'doc']:
                        file_info, error = extract_docx_detailed(filepath)
                    elif file_ext == 'txt':
                        file_info, error = extract_txt_detailed(filepath)
                    else:
                        error = "Unsupported file type"
                        file_info = None
                   
                    if error:
                        st.error(f"❌ {filename}: {error}")
                        continue
                   
                    save_cache(cache_key, file_info)
                    st.success(f"💾 Cached: {filename}")
               
                files_data[filename] = file_info
               
                # Add chunks with full metadata
                for chunk_obj in file_info['chunks']:
                    # Handle both dict and string formats
                    if isinstance(chunk_obj, dict):
                        all_chunks.append(chunk_obj['content'])
                        all_metadata.append(chunk_obj['metadata'])
                    else:
                        # Fallback for old cached data
                        all_chunks.append(chunk_obj)
                        all_metadata.append({
                            "source": filename, 
                            "page": "N/A",
                            "is_table": "False",
                            "table_number": "N/A"
                        })
               
                progress_bar.progress((idx + 1) / len(available_files))
           
            status_text.text("Building search index...")
           
            if all_chunks:
                batch_size = 500
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i+batch_size]
                    metadata_batch = all_metadata[i:i+batch_size]
                    collection.add(
                        documents=batch,
                        ids=[f"chunk_{i+j}" for j in range(len(batch))],
                        metadatas=metadata_batch
                    )
           
            st.session_state.files_data = files_data
            st.session_state.collection = collection
            st.session_state.processed = True
            st.session_state.messages = []  # Reset chat
           
            status_text.empty()
            st.success("✅ Processing completed!")
            st.balloons()
    
    if st.session_state.processed:
        st.markdown("---")
        st.markdown("### 📊 Statistics")
        total_chunks = sum(len(info['chunks']) for info in st.session_state.files_data.values())
        total_tables = sum(info['total_tables'] for info in st.session_state.files_data.values())
        
        st.metric("Files", len(st.session_state.files_data))
        st.metric("Chunks", total_chunks)
        st.metric("Tables", total_tables)
        
        st.markdown("---")
        
        if st.button("🔄 Clear Chat History"):
            st.session_state.messages = []
            st.session_state.current_context = []
            st.rerun()
        
        if st.button("🗑️ Clear Cache & Reprocess"):
            import shutil
            if os.path.exists(CACHE_FOLDER):
                shutil.rmtree(CACHE_FOLDER)
                os.makedirs(CACHE_FOLDER)
            st.session_state.processed = False
            st.session_state.files_data = {}
            st.session_state.collection = None
            st.session_state.messages = []
            st.success("✅ Cache cleared! Click 'Process Documents' to reprocess.")
            st.rerun()

# Main chat interface
if st.session_state.processed:
    st.markdown("### 💬 Chat with Documents")
    
    # Display chat history
    for message in st.session_state.messages:
        role = message["role"]
        content = message["content"]
        
        if role == "user":
            st.markdown(f'<div class="chat-message user-message">👤 <b>You:</b> {content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-message assistant-message">🤖 <b>Assistant:</b><br>{content}</div>', unsafe_allow_html=True)
    
    # Chat input
    query = st.chat_input("Ask anything about your documents...")
    
    if query:
        # Add user message
        st.session_state.messages.append({"role": "user", "content": query})
        
        with st.spinner("Thinking..."):
            # Search with metadata
            results = st.session_state.collection.query(
                query_texts=[query],
                n_results=10
            )
            
            # Build chunk objects with metadata
            relevant_chunks = []
            for content, metadata in zip(results["documents"][0], results["metadatas"][0]):
                relevant_chunks.append({
                    'content': content,
                    'metadata': metadata
                })
            
            # Generate answer with chat history
            answer = answer_question_with_groq(query, relevant_chunks, st.session_state.messages)
            
            # Add assistant message
            st.session_state.messages.append({"role": "assistant", "content": answer})
            
            # Store context for follow-ups
            st.session_state.current_context = relevant_chunks
        
        st.rerun()
    
    # Show sources in expander
    if st.session_state.current_context:
        with st.expander("📄 View Sources", expanded=False):
            for idx, chunk_data in enumerate(st.session_state.current_context[:5], 1):
                meta = chunk_data['metadata']
                
                # Handle string metadata
                source = meta.get('source', 'Unknown')
                page = meta.get('page', 'N/A')
                is_table = meta.get('is_table', 'False')
                table_num = meta.get('table_number', 'N/A')
                
                citation_info = f"📄 **Source {idx}**: {source} | Page {page}"
                if is_table == 'True' or is_table == True:
                    citation_info += f" | Table {table_num}"
                
                st.markdown(citation_info)
                st.markdown(f'<div class="chunk-display">{chunk_data["content"][:500]}...</div>', unsafe_allow_html=True)
                st.markdown("---")

else:
    st.info("👈 Click 'Process Documents' in the sidebar to start!")
    
    st.markdown("""
    ### 🎯 Features:
    - **Precise Citations**: Every answer includes file + page + table references
    - **Conversational**: Ask follow-up questions naturally ("summarize that", "tell me more")
    - **Smart Context**: Understands when you refer to previous answers
    - **Multi-language**: Supports English, German, and Arabic
    - **Fast**: Cached processing for instant responses
    - **MBE-Specific**: Optimized for biomedical engineering regulations
    
    ### 📋 Supported Documents:
    - 📄 Study & Examination Regulations (SPO)
    - 📚 Module Handbook
    - 📝 Guide for Writing Scientific Papers
    - 📃 Notes on Bachelor/Master Theses
    - ✍️ Scientific Writing Guidelines
    
    ### 💡 Example Questions:
    - "How many modules in semester 1?"
    - "What are the thesis requirements?"
    - "Tell me about the internship" → then "summarize that"
    - "Compare exam types in SPO"
    """)
